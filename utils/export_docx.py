from docx import Document


def create_docx(text, filename):
    """
    Creates a Word document from text.

    Parameters:
        text (str): Content to write.
        filename (str): Output DOCX filename.
    """

    document = Document()

    document.add_heading("AI Study Assistant", level=1)

    for line in text.split("\n"):
        document.add_paragraph(line)

    document.save(filename)

    return filename